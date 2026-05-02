"""
Скрипт парсинга документов Word (.doc, .docx) и PDF.
Извлекает текст → сохраняет в data/cleaned/docs/ в формате JSON,
совместимом с остальным пайплайном (url, title, content).

Запуск:
    python scripts/02_parse_docs.py
"""

import json
import os
import re
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

# Добавляем корень проекта в sys.path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

DOCS_DIR = PROJECT_ROOT / "data" / "raw" / "docs"
OUTPUT_DIR = PROJECT_ROOT / "data" / "cleaned" / "docs"


# ── Парсеры по типам файлов ──────────────────────────────────────────


def parse_docx(filepath: Path) -> str:
    """Извлекает текст из .docx файла."""
    from docx import Document

    doc = Document(str(filepath))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return "\n".join(paragraphs)


def parse_doc(filepath: Path) -> str:
    """
    Извлекает текст из .doc (старый бинарный формат Word).
    Стратегия:
      1. COM-автоматизация через Microsoft Word (pywin32)
      2. Конвертация через LibreOffice (headless)
    """
    # Стратегия 1: COM-автоматизация через Word (Windows + MS Word)
    try:
        import win32com.client
        import pythoncom

        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(filepath.resolve()))
            text = doc.Content.Text
            doc.Close(False)
            if text and text.strip():
                print("  [OK] Извлечено через Word COM")
                return text.strip()
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
    except ImportError:
        print("  [WARN] pywin32 не установлен, пропускаем COM")
    except Exception as e:
        print(f"  [WARN] Word COM не удался: {e}")

    # Стратегия 2: конвертация в .docx через LibreOffice
    try:
        import subprocess
        import tempfile
        import shutil

        lo_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice",
        ]
        soffice = None
        for p in lo_paths:
            if os.path.isfile(p) or shutil.which(p):
                soffice = p
                break

        if soffice:
            with tempfile.TemporaryDirectory() as tmpdir:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "docx",
                     "--outdir", tmpdir, str(filepath)],
                    capture_output=True, timeout=60,
                )
                converted = Path(tmpdir) / (filepath.stem + ".docx")
                if converted.exists():
                    print("  [OK] Извлечено через LibreOffice")
                    return parse_docx(converted)
    except Exception as e:
        print(f"  [WARN] LibreOffice конвертация не удалась: {e}")

    print(f"  [ERROR] Не удалось извлечь текст из .doc: {filepath.name}")
    print("         Установите Microsoft Word или LibreOffice")
    return ""


def parse_pdf(filepath: Path) -> str:
    """Извлекает текст из PDF файла."""
    import pdfplumber

    pages_text = []
    with pdfplumber.open(str(filepath)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text.strip())

    return "\n".join(pages_text)


# ── Утилиты ──────────────────────────────────────────────────────────


def extract_title(filename: str) -> str:
    """Извлекает читаемый заголовок из имени файла."""
    name = Path(filename).stem
    # Заменяем _ЫЫ на пробел + ЫЫ (инициалы)
    name = re.sub(r"_([А-ЯA-Z]{1,2})$", r" \1", name)
    name = re.sub(r"_([А-ЯA-Z]{1,2})_", r" \1 ", name)
    # Остальные подчёркивания → пробелы
    name = name.replace("_", " ")
    return name.strip()


def clean_extracted_text(text: str) -> str:
    """Базовая очистка извлечённого текста."""
    # Убираем множественные пустые строки
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Убираем множественные пробелы
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


# ── Основная логика ──────────────────────────────────────────────────


PARSERS = {
    ".docx": parse_docx,
    ".doc": parse_doc,
    ".pdf": parse_pdf,
}

SUPPORTED_EXTENSIONS = set(PARSERS.keys())


def main():
    if not DOCS_DIR.exists():
        print(f"Директория не найдена: {DOCS_DIR}")
        return

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    files = [
        f for f in sorted(DOCS_DIR.iterdir())
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    if not files:
        print(f"Не найдено документов (.doc, .docx, .pdf) в {DOCS_DIR}")
        return

    print(f"Найдено документов: {len(files)}")
    print("=" * 70)

    stats = []
    for filepath in files:
        ext = filepath.suffix.lower()
        parser = PARSERS[ext]
        title = extract_title(filepath.name)

        print(f"\n[{ext:5}] {filepath.name}")

        try:
            raw_text = parser(filepath)
            if not raw_text:
                print(f"  [SKIP] Пустой текст, пропускаем")
                continue

            content = clean_extracted_text(raw_text)

            doc_json = {
                "url": f"local://docs/{filepath.name}",
                "title": title,
                "content": content,
            }

            out_filename = filepath.stem + ".json"
            out_path = OUTPUT_DIR / out_filename

            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(doc_json, f, ensure_ascii=False, indent=2)

            stats.append({
                "file": filepath.name,
                "title": title,
                "content_len": len(content),
                "output": out_path.name,
            })

            print(f"  title: {title}")
            print(f"  content: {len(content)} символов")
            print(f"  -> {out_path.name}")

        except Exception as e:
            print(f"  [ERROR] {e}")

    print(f"\n{'=' * 70}")
    print(f"Обработано файлов: {len(stats)} / {len(files)}")
    total_chars = sum(s["content_len"] for s in stats)
    print(f"Общий объём текста: {total_chars} символов")
    print(f"Результат в: {OUTPUT_DIR.resolve()}")


if __name__ == "__main__":
    main()
